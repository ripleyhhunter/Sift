"""
Document processing module for Sift.

Handles text extraction from documents (PDF, Office, CSV) for fast text-based classification,
with fallback to image conversion for vision model analysis of image files.
"""

import io
import os
import csv
import base64
import shutil
import logging
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple, Union
from PIL import Image

from .config import Config
from .utils import ensure_directory
from .platform_utils import get_subprocess_flags, get_libreoffice_path, IS_WINDOWS, IS_MACOS

logger = logging.getLogger(__name__)

# Image file extensions (require vision model)
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp'}

# Text-extractable extensions (can use fast text model)
PDF_EXTENSION = '.pdf'
WORD_EXTENSIONS = {'.docx', '.doc'}
EXCEL_EXTENSIONS = {'.xlsx', '.xls'}
CSV_EXTENSIONS = {'.csv', '.tsv'}
POWERPOINT_EXTENSIONS = {'.pptx', '.ppt'}

# All Office extensions
OFFICE_EXTENSIONS = WORD_EXTENSIONS | EXCEL_EXTENSIONS | POWERPOINT_EXTENSIONS | {'.odt', '.ods', '.odp'}

# Extensions that support text extraction
TEXT_EXTRACTABLE = {PDF_EXTENSION} | WORD_EXTENSIONS | EXCEL_EXTENSIONS | CSV_EXTENSIONS


class DocumentProcessor:
    """
    Process documents for classification.
    
    Supports two modes:
    1. Text extraction (fast) - For PDF, DOCX, XLSX, CSV
    2. Image conversion (slower) - For image files or when text extraction fails
    """
    
    def __init__(self, config: Config):
        """
        Initialize the document processor.
        
        Args:
            config: Application configuration
        """
        self.config = config
        self.temp_path = config.folders.temp_path
        self.max_pages = config.processing.max_pages_to_analyze
        self.image_dpi = config.processing.image_dpi
        self.max_dimension = config.processing.max_image_dimension
        self.max_text_chars = 8000  # Limit text to avoid token limits
        
        # Ensure temp directory exists
        ensure_directory(self.temp_path)
        
        # Check for required tools
        self._poppler_available = self._check_poppler()
        self._libreoffice_available = self._check_libreoffice()
        self._text_extractors_available = self._check_text_extractors()
        self._tesseract_available = self._check_tesseract()
        
        # OCR settings
        self._ocr_enabled = config.processing.ocr_enabled if hasattr(config.processing, 'ocr_enabled') else True
        self._ocr_min_text_threshold = 50  # If less than N chars extracted, try OCR
    
    def _check_text_extractors(self) -> dict:
        """Check which text extraction libraries are available."""
        available = {}
        
        try:
            import pypdf
            available['pdf'] = True
        except ImportError:
            available['pdf'] = False
            logger.debug("pypdf not available for PDF text extraction")
        
        try:
            import docx
            available['docx'] = True
        except ImportError:
            available['docx'] = False
            logger.debug("python-docx not available for DOCX text extraction")
        
        try:
            import openpyxl
            available['xlsx'] = True
        except ImportError:
            available['xlsx'] = False
            logger.debug("openpyxl not available for XLSX text extraction")
        
        # CSV is built-in, always available
        available['csv'] = True
        
        return available
    
    def _check_poppler(self) -> bool:
        """Check if Poppler (pdf2image dependency) is available."""
        try:
            result = subprocess.run(
                ['pdftoppm', '-v'],
                capture_output=True,
                text=True,
                **get_subprocess_flags()
            )
            return True
        except FileNotFoundError:
            logger.debug("Poppler not found - PDF image conversion unavailable")
            return False
        except Exception as e:
            logger.debug(f"Error checking Poppler: {e}")
            return False
    
    def _check_libreoffice(self) -> bool:
        """Check if LibreOffice is available by checking if executable exists."""
        try:
            soffice_path = get_libreoffice_path()
            if soffice_path:
                self._soffice_path = soffice_path
                logger.debug(f"Found LibreOffice at: {soffice_path}")
                return True
            logger.debug("LibreOffice not found in standard locations")
            return False
        except (OSError, PermissionError) as e:
            logger.debug(f"Error checking for LibreOffice: {e}")
            return False
    
    def _check_tesseract(self) -> bool:
        """Check if Tesseract OCR is available."""
        try:
            result = subprocess.run(
                ['tesseract', '--version'],
                capture_output=True,
                text=True,
                **get_subprocess_flags()
            )
            if result.returncode == 0:
                version = result.stdout.split('\n')[0] if result.stdout else 'unknown'
                logger.debug(f"Tesseract OCR available: {version}")
                return True
            return False
        except FileNotFoundError:
            logger.debug("Tesseract OCR not found - OCR extraction unavailable")
            return False
        except Exception as e:
            logger.debug(f"Error checking Tesseract: {e}")
            return False
    
    @property
    def ocr_available(self) -> bool:
        """Check if OCR capability is available."""
        return self._tesseract_available and self._ocr_enabled
    
    def can_extract_text(self, file_path: Path) -> bool:
        """
        Check if text can be extracted from this file type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if text extraction is possible
        """
        suffix = file_path.suffix.lower()
        
        if suffix == PDF_EXTENSION:
            return self._text_extractors_available.get('pdf', False)
        elif suffix in WORD_EXTENSIONS:
            return self._text_extractors_available.get('docx', False)
        elif suffix in EXCEL_EXTENSIONS:
            return self._text_extractors_available.get('xlsx', False)
        elif suffix in CSV_EXTENSIONS:
            return True  # Built-in
        
        return False
    
    def is_image_file(self, file_path: Path) -> bool:
        """Check if file is an image that requires vision model."""
        return file_path.suffix.lower() in IMAGE_EXTENSIONS
    
    def extract_text(self, file_path: Path) -> Optional[str]:
        """
        Extract text content from a document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text, or None if extraction failed
        """
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == PDF_EXTENSION:
                text = self._extract_pdf_text(file_path)
            elif suffix == '.docx':
                text = self._extract_docx_text(file_path)
            elif suffix == '.doc':
                text = self._extract_doc_text(file_path)
            elif suffix in {'.xlsx', '.xls'}:
                text = self._extract_excel_text(file_path)
            elif suffix in CSV_EXTENSIONS:
                text = self._extract_csv_text(file_path)
            else:
                logger.debug(f"No text extractor for {suffix}")
                return None
            
            if text:
                # Clean and limit text
                text = self._clean_text(text)
                if len(text) > self.max_text_chars:
                    text = text[:self.max_text_chars] + "\n\n[... content truncated ...]"
                
                logger.debug(f"Extracted {len(text)} characters from {file_path.name}")
                return text
            
            return None
            
        except Exception as e:
            logger.warning(f"Text extraction failed for {file_path.name}: {e}")
            return None
    
    def _extract_pdf_text(self, file_path: Path) -> Optional[str]:
        """
        Extract text from PDF using pypdf, with OCR fallback for scanned documents.
        
        If the PDF appears to be scanned (minimal extractable text), falls back
        to OCR using Tesseract if available.
        """
        try:
            import pypdf
            
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                max_pages = min(len(reader.pages), self.max_pages)
                
                for i in range(max_pages):
                    page = reader.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {i+1} ---\n{page_text}")
            
            combined_text = "\n\n".join(text_parts) if text_parts else ""
            
            # Check if this appears to be a scanned PDF (very little text extracted)
            clean_text = combined_text.replace(" ", "").replace("\n", "")
            if len(clean_text) < self._ocr_min_text_threshold:
                logger.debug(f"PDF has minimal text ({len(clean_text)} chars), attempting OCR")
                ocr_text = self._extract_pdf_ocr(file_path)
                if ocr_text and len(ocr_text) > len(combined_text):
                    logger.info(f"OCR extracted {len(ocr_text)} chars from scanned PDF")
                    return ocr_text
            
            return combined_text if combined_text else None
            
        except Exception as e:
            logger.debug(f"PDF text extraction error: {e}")
            # Try OCR as last resort
            if self.ocr_available:
                return self._extract_pdf_ocr(file_path)
            return None
    
    def _extract_pdf_ocr(self, file_path: Path) -> Optional[str]:
        """
        Extract text from a PDF using OCR (for scanned documents).
        
        Converts PDF pages to images, then runs Tesseract on each.
        """
        if not self.ocr_available:
            return None
        
        if not self._poppler_available:
            logger.debug("PDF OCR requires Poppler for image conversion")
            return None
        
        try:
            from pdf2image import convert_from_path
            
            # Convert PDF pages to images
            images = convert_from_path(
                file_path,
                first_page=1,
                last_page=min(self.max_pages, 3),  # Limit pages for OCR (expensive)
                dpi=150  # Lower DPI for OCR (faster, still readable)
            )
            
            text_parts = []
            for i, image in enumerate(images):
                page_text = self._ocr_image(image)
                if page_text:
                    text_parts.append(f"--- Page {i+1} (OCR) ---\n{page_text}")
            
            return "\n\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.debug(f"PDF OCR extraction error: {e}")
            return None
    
    def _ocr_image(self, image) -> Optional[str]:
        """
        Run OCR on a PIL Image using Tesseract.
        
        Args:
            image: PIL Image object
            
        Returns:
            Extracted text or None
        """
        if not self._tesseract_available:
            return None
        
        try:
            import pytesseract
            
            # Run OCR with English language
            text = pytesseract.image_to_string(
                image,
                lang='eng',
                config='--psm 1'  # Automatic page segmentation with OSD
            )
            
            return text.strip() if text else None
            
        except ImportError:
            # pytesseract not installed, try direct subprocess
            return self._ocr_image_subprocess(image)
        except Exception as e:
            logger.debug(f"OCR error: {e}")
            return None
    
    def _ocr_image_subprocess(self, image) -> Optional[str]:
        """Run OCR using tesseract subprocess directly."""
        try:
            # Save image to temp file
            temp_path = self.temp_path / f"ocr_temp_{os.getpid()}.png"
            image.save(str(temp_path), 'PNG')
            
            try:
                # Run tesseract
                result = subprocess.run(
                    ['tesseract', str(temp_path), 'stdout', '-l', 'eng'],
                    capture_output=True,
                    text=True,
                    timeout=30,
                    **get_subprocess_flags()
                )
                
                if result.returncode == 0:
                    return result.stdout.strip()
                return None
                
            finally:
                # Clean up temp file
                if temp_path.exists():
                    temp_path.unlink()
                    
        except Exception as e:
            logger.debug(f"Subprocess OCR error: {e}")
            return None
    
    def _extract_docx_text(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX using python-docx."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.debug(f"DOCX text extraction error: {e}")
            return None
    
    def _extract_doc_text(self, file_path: Path) -> Optional[str]:
        """Extract text from legacy DOC files (via LibreOffice conversion)."""
        if not self._libreoffice_available:
            return None
        
        # Convert to DOCX first, then extract
        with tempfile.TemporaryDirectory(dir=str(self.temp_path)) as temp_dir:
            try:
                cmd = [
                    self._soffice_path,
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', temp_dir,
                    str(file_path)
                ]
                
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=30,
                    **get_subprocess_flags()
                )
                
                if result.returncode == 0:
                    docx_path = Path(temp_dir) / (file_path.stem + '.docx')
                    if docx_path.exists():
                        return self._extract_docx_text(docx_path)
                
            except Exception as e:
                logger.debug(f"DOC conversion error: {e}")
        
        return None
    
    def _extract_excel_text(self, file_path: Path) -> Optional[str]:
        """Extract text from Excel files using openpyxl."""
        try:
            import openpyxl
            
            suffix = file_path.suffix.lower()
            
            # Handle .xls files by converting first
            if suffix == '.xls':
                return self._extract_xls_text(file_path)
            
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames[:3]:  # Limit to first 3 sheets
                sheet = wb[sheet_name]
                text_parts.append(f"--- Sheet: {sheet_name} ---")
                
                row_count = 0
                for row in sheet.iter_rows(max_row=100, values_only=True):  # Limit rows
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        text_parts.append(row_text)
                        row_count += 1
                        if row_count >= 50:  # Limit rows per sheet
                            text_parts.append("[... more rows ...]")
                            break
            
            wb.close()
            return "\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.debug(f"Excel text extraction error: {e}")
            return None
    
    def _extract_xls_text(self, file_path: Path) -> Optional[str]:
        """Extract text from legacy XLS files (via LibreOffice conversion)."""
        if not self._libreoffice_available:
            return None
        
        with tempfile.TemporaryDirectory(dir=str(self.temp_path)) as temp_dir:
            try:
                cmd = [
                    self._soffice_path,
                    '--headless',
                    '--convert-to', 'xlsx',
                    '--outdir', temp_dir,
                    str(file_path)
                ]
                
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=30,
                    **get_subprocess_flags()
                )
                
                if result.returncode == 0:
                    xlsx_path = Path(temp_dir) / (file_path.stem + '.xlsx')
                    if xlsx_path.exists():
                        return self._extract_excel_text(xlsx_path)
                
            except Exception as e:
                logger.debug(f"XLS conversion error: {e}")
        
        return None
    
    def _extract_csv_text(self, file_path: Path) -> Optional[str]:
        """Extract text from CSV/TSV files."""
        try:
            delimiter = '\t' if file_path.suffix.lower() == '.tsv' else ','
            text_parts = []
            
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding, newline='') as f:
                        reader = csv.reader(f, delimiter=delimiter)
                        row_count = 0
                        for row in reader:
                            row_text = " | ".join(str(cell) for cell in row if cell)
                            if row_text.strip():
                                text_parts.append(row_text)
                                row_count += 1
                                if row_count >= 100:  # Limit rows
                                    text_parts.append("[... more rows ...]")
                                    break
                        break  # Success, stop trying encodings
                except UnicodeDecodeError:
                    continue
            
            return "\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.debug(f"CSV text extraction error: {e}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        empty_count = 0
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
                empty_count = 0
            else:
                empty_count += 1
                if empty_count <= 2:  # Allow max 2 consecutive empty lines
                    cleaned_lines.append("")
        
        return "\n".join(cleaned_lines)
    
    # ========== Image processing methods (for vision model fallback) ==========
    
    def process_document(self, file_path: Path) -> List[str]:
        """
        Convert document to list of base64-encoded images (for vision model).
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of base64-encoded PNG images
        """
        suffix = file_path.suffix.lower()
        
        try:
            if suffix in IMAGE_EXTENSIONS:
                images = self._process_image(file_path)
            elif suffix == PDF_EXTENSION:
                images = self._process_pdf(file_path)
            elif suffix in OFFICE_EXTENSIONS:
                images = self._process_office(file_path)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
            
            base64_images = []
            for img in images:
                optimized = self._optimize_image(img)
                b64 = self._image_to_base64(optimized)
                base64_images.append(b64)
            
            logger.debug(f"Processed {len(base64_images)} page(s) from {file_path.name}")
            return base64_images
            
        except Exception as e:
            logger.error(f"Error processing document {file_path.name}: {e}")
            raise DocumentProcessingError(f"Failed to process {file_path.name}: {e}")
        finally:
            self._cleanup_temp_files()
    
    def _process_image(self, file_path: Path) -> List[Image.Image]:
        """Process an image file."""
        try:
            img = Image.open(file_path)
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if 'A' in img.mode else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            return [img]
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process image: {e}")
    
    def _process_pdf(self, file_path: Path) -> List[Image.Image]:
        """Convert PDF pages to images."""
        if not self._poppler_available:
            raise DocumentProcessingError("Poppler not available for PDF image conversion")
        
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(
                file_path,
                dpi=self.image_dpi,
                first_page=1,
                last_page=self.max_pages,
                fmt='png',
                thread_count=2
            )
            return list(images)
        except Exception as e:
            raise DocumentProcessingError(f"Failed to convert PDF: {e}")
    
    def _process_office(self, file_path: Path) -> List[Image.Image]:
        """Convert Office document to images via PDF."""
        if not self._libreoffice_available:
            raise DocumentProcessingError("LibreOffice not available")
        
        with tempfile.TemporaryDirectory(dir=str(self.temp_path)) as temp_dir:
            try:
                pdf_path = self._convert_office_to_pdf(file_path, Path(temp_dir))
                if not pdf_path or not pdf_path.exists():
                    raise DocumentProcessingError(f"Conversion failed for {file_path.name}")
                return self._process_pdf(pdf_path)
            except Exception as e:
                raise DocumentProcessingError(f"Failed to process Office document: {e}")
    
    def _convert_office_to_pdf(self, input_path: Path, output_dir: Path) -> Optional[Path]:
        """Convert Office document to PDF using LibreOffice."""
        try:
            cmd = [
                self._soffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                str(input_path)
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
                **get_subprocess_flags()
            )
            
            if result.returncode != 0:
                logger.debug(f"LibreOffice conversion failed: {result.stderr}")
                return None
            
            pdf_path = output_dir / (input_path.stem + '.pdf')
            if pdf_path.exists():
                return pdf_path
            
            pdfs = list(output_dir.glob('*.pdf'))
            return pdfs[0] if pdfs else None
            
        except subprocess.TimeoutExpired:
            logger.warning(f"LibreOffice conversion timed out for {input_path.name}")
            return None
        except (OSError, subprocess.SubprocessError) as e:
            logger.debug(f"LibreOffice conversion error: {e}")
            return None
    
    def _optimize_image(self, image: Image.Image) -> Image.Image:
        """Resize and optimize image for LLM processing."""
        width, height = image.size
        if width <= self.max_dimension and height <= self.max_dimension:
            return image
        
        if width > height:
            new_width = self.max_dimension
            new_height = int(height * (self.max_dimension / width))
        else:
            new_height = self.max_dimension
            new_width = int(width * (self.max_dimension / height))
        
        return image.resize((new_width, new_height), Image.Resampling.LANCZOS)
    
    def _image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string."""
        buffer = io.BytesIO()
        image.save(buffer, format='PNG', optimize=True)
        buffer.seek(0)
        return base64.b64encode(buffer.read()).decode('utf-8')
    
    def _cleanup_temp_files(self) -> None:
        """Remove temporary conversion files."""
        try:
            if self.temp_path.exists():
                for item in self.temp_path.iterdir():
                    try:
                        if item.is_file():
                            item.unlink()
                        elif item.is_dir():
                            shutil.rmtree(item)
                    except PermissionError:
                        # File may still be in use, will be cleaned up next time
                        logger.debug(f"Could not remove temp file (in use): {item}")
                    except OSError as e:
                        logger.debug(f"Could not remove temp item {item}: {e}")
        except OSError as e:
            logger.debug(f"Could not access temp directory: {e}")
    
    def get_capabilities(self) -> dict:
        """Get information about processing capabilities."""
        return {
            'images': True,
            'pdf': self._poppler_available,
            'office': self._libreoffice_available,
            'text_extraction': {
                'pdf': self._text_extractors_available.get('pdf', False),
                'docx': self._text_extractors_available.get('docx', False),
                'xlsx': self._text_extractors_available.get('xlsx', False),
                'csv': True
            }
        }


class DocumentProcessingError(Exception):
    """Exception raised for document processing errors."""
    pass
